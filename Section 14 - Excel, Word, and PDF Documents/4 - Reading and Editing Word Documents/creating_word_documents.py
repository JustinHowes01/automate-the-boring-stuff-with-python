import docx

doc = docx.Document()

doc.add_paragraph('Hello, this is a paragraph. ')
doc.add_paragraph('This is another paragraph. ')

p = doc.paragraphs[0]
p.add_run('This is a new run. ')
p.runs[1].bold = True

doc.save('demo3.docx')