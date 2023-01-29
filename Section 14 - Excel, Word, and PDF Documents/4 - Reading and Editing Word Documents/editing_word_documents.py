import docx

doc = docx.Document('demo.docx')

print(doc.paragraphs[0].text, end='\n\n')

p = doc.paragraphs[1]

print(p.runs, end='\n\n')

print(''.join([i.text for i in p.runs]), end='\n\n')

p.runs[3].underline = True
p.runs[3].text = 'italic and underline'

doc.save('demo2.docx')